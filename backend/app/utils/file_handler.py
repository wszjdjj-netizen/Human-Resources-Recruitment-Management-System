"""
文件处理工具：保存上传文件、提取PDF/Word文本
支持扫描版PDF的OCR降级检测
"""
import logging
import os
from pathlib import Path
from app.config import get_settings

logger = logging.getLogger(__name__)

settings = get_settings()

# 支持的文件类型
ALLOWED_EXTENSIONS = {".pdf", ".docx"}


class TextExtractionResult:
    """文本提取结果，携带元信息"""
    def __init__(self, text: str, is_empty: bool = False, reason: str = "", method: str = ""):
        self.text = text or ""
        self.is_empty = is_empty
        self.reason = reason  # 为空时的原因说明
        self.method = method     # 使用的提取方法

    @property
    def ok(self) -> bool:
        return not self.is_empty and len(self.text.strip()) >= 10


def get_upload_dir() -> Path:
    """获取上传文件目录（确保目录存在）"""
    upload_dir = Path(settings.upload_dir)
    if not upload_dir.is_absolute():
        upload_dir = Path.cwd() / upload_dir
    upload_dir.mkdir(parents=True, exist_ok=True)
    return upload_dir


def sanitize_filename(filename: str) -> str:
    """清理文件名，防止目录遍历攻击"""
    import re
    basename = Path(filename).name
    cleaned = re.sub(r'[^a-zA-Z0-9._-]', '_', basename)
    if not cleaned or cleaned.startswith('.'):
        cleaned = 'file_' + cleaned
    return cleaned[:255]


def save_upload_file(file_content: bytes, filename: str) -> str:
    """保存上传文件到 uploads 目录，返回完整路径"""
    upload_dir = get_upload_dir()
    cleaned_filename = sanitize_filename(filename)
    import time
    safe_name = f"{int(time.time() * 1000)}_{cleaned_filename}"
    file_path = upload_dir / safe_name
    if not file_path.resolve().is_relative_to(upload_dir.resolve()):
        raise ValueError("非法的文件路径")
    file_path.write_bytes(file_content)
    return str(file_path)


def delete_uploaded_temp_file(file_path: str) -> None:
    """删除后端 uploads 目录中的临时上传副本。"""
    upload_dir = get_upload_dir().resolve()
    target = Path(file_path).resolve()
    if not target.is_relative_to(upload_dir):
        raise ValueError("拒绝删除上传目录之外的文件")
    target.unlink(missing_ok=True)


def extract_text_from_pdf(file_path: str) -> TextExtractionResult:
    """从PDF文件中提取文本，自动检测扫描版PDF"""
    from PyPDF2 import PdfReader
    try:
        reader = PdfReader(file_path)
    except Exception as e:
        logger.warning(f"无法读取PDF文件 {file_path}: {e}")
        return TextExtractionResult("", is_empty=True, reason=f"PDF文件损坏或无法读取: {e}", method="pydf2")

    text_parts = []
    total_pages = len(reader.pages)
    for page_idx, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text)
        except Exception as e:
            logger.warning(f"PDF第{page_idx+1}页提取失败: {e}")

    full_text = "\n".join(text_parts).strip()
    # 检测是否为空（很可能是扫描版/图片型PDF）
    if not full_text or len(full_text) < 10:
        # 进一步检查：如果PDF有页面但无文字，几乎可以确定是扫描版
        if total_pages > 0:
            reason = f"PDF共{total_pages}页但未提取到文字内容，可能是扫描版/图片型PDF"
            # 尝试用 pdfplumber 做二次提取（对某些加密PDF效果更好）
            try:
                from pdfplumber import open as pdf_open, PdfReader as PlumberReader
                with pdf_open(file_path) as pdf:
                    pt = []
                    for p in pdf.pages:
                        t = p.extract_text()
                        if t: pt.append(t)
                    alt_text = "\n".join(pt).strip()
                    if alt_text and len(alt_text) >= 10:
                        logger.info(f"pdfplumber成功提取到 {len(alt_text)} 字符")
                        return TextExtractionResult(alt_text, is_empty=False, method="pdfplumber")
            except ImportError:
                pass
            except Exception as e2:
                logger.debug(f"pdfplumber也失败了: {e2}")
            return TextExtractionResult(full_text, is_empty=True, reason=reason, method="pydf2")
    return TextExtractionResult(full_text, is_empty=False, method="pydf2")


def extract_text_from_docx(file_path: str) -> TextExtractionResult:
    """从Word(.docx)文件中提取文本"""
    try:
        from docx import Document
        doc = Document(file_path)
    except Exception as e:
        return TextExtractionResult("", is_empty=True, reason=f"Word文件读取失败: {e}", method="python-docx")

    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    full_text = "\n".join(text_parts).strip()
    if not full_text or len(full_text) < 5:
        return TextExtractionResult(full_text, is_empty=True, reason="Word文档内容为空或过短", method="python-docx")
    return TextExtractionResult(full_text, is_empty=False, method="python-docx")


def extract_text(file_path: str) -> TextExtractionResult:
    """
    根据文件扩展名自动选择提取方法，返回TextExtractionResult对象
    兼容旧接口：.text 属性返回原始文本字符串
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"不支持的文件格式：{ext}，请上传 PDF 或 Word 文件")


# ===== 以下为兼容旧接口的便捷函数 =====

def extract_text_raw(file_path: str) -> str:
    """兼容旧接口：返回纯文本字符串（建议使用extract_text()获取更详细结果）"""
    result = extract_text(file_path)
    return result.text


def is_allowed_file(filename: str) -> bool:
    """检查文件扩展名是否允许"""
    ext = Path(filename).suffix.lower()
    return ext in ALLOWED_EXTENSIONS
